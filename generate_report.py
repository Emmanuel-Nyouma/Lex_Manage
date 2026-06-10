"""
LexManage BTech SWE Report Generator
Produces a fully formatted .docx with Times New Roman, proper heading styles,
and screenshot placeholders for every section / sub-section.
"""
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page setup ───────────────────────────────────────────────────────────────
section = doc.sections[0]
section.page_width  = Cm(21)
section.page_height = Cm(29.7)
section.top_margin    = Cm(2.54)
section.bottom_margin = Cm(2.54)
section.left_margin   = Cm(3.0)
section.right_margin  = Cm(2.5)

TNR = "Times New Roman"

# ── Helper utilities ─────────────────────────────────────────────────────────
def set_font(run, size=12, bold=False, italic=False, color=None):
    run.font.name = TNR
    run.font.size = Pt(size)
    run.font.bold  = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)
    # Force Times New Roman via XML theme override
    r = run._r
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), TNR)
    rFonts.set(qn('w:hAnsi'), TNR)
    rFonts.set(qn('w:cs'), TNR)
    existing = rPr.find(qn('w:rFonts'))
    if existing is not None:
        rPr.remove(existing)
    rPr.insert(0, rFonts)


def para(text, size=12, bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
         space_before=0, space_after=6, color=None):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    p.paragraph_format.line_spacing = Pt(22)  # ~1.5 line spacing
    run = p.add_run(text)
    set_font(run, size, bold, italic, color)
    return p


def heading(text, level=1, size=None):
    """Chapter headings (CHAPTER I, CHAPTER II, …) or section headings."""
    sizes = {1: 14, 2: 13, 3: 12}
    s = size or sizes.get(level, 12)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 0 else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    set_font(run, s, bold=True)
    return p


def chapter_title(num_str, title_text):
    """E.g. chapter_title('CHAPTER I', 'GENERAL INTRODUCTION')"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after  = Pt(0)
    run = p.add_run(num_str)
    set_font(run, 14, bold=True)
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after  = Pt(18)
    r2 = p2.add_run(title_text)
    set_font(r2, 14, bold=True)


def section_heading(number, title, size=12):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(f"{number} {title}")
    set_font(run, size, bold=True)
    return p


def screenshot_box(label):
    """Inserts a labelled placeholder box for a screenshot."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(4)

    # Draw a simple bordered paragraph via XML shading
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    for side in ('top', 'left', 'bottom', 'right'):
        bd = OxmlElement(f'w:{side}')
        bd.set(qn('w:val'), 'single')
        bd.set(qn('w:sz'), '12')
        bd.set(qn('w:space'), '4')
        bd.set(qn('w:color'), 'B0B0B0')
        pBdr.append(bd)
    pPr.append(pBdr)

    run = p.add_run(f"[ SCREENSHOT PLACEHOLDER — {label} ]")
    set_font(run, 10, italic=True, color=(128, 128, 128))

    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_before = Pt(2)
    caption.paragraph_format.space_after  = Pt(12)
    cr = caption.add_run(f"Figure X: {label}")
    set_font(cr, 10, italic=True)


def bullet(text, size=12):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(3)
    run = p.add_run(text)
    set_font(run, size)
    return p


def page_break():
    doc.add_page_break()


# ─────────────────────────────────────────────────────────────────────────────
# FRONT MATTER
# ─────────────────────────────────────────────────────────────────────────────

# ── ATTESTATION ──────────────────────────────────────────────────────────────
heading("ATTESTATION", level=0)
para("")
para(
    "I, the undersigned, Emmanuel NYOUMA, hereby attest that this work is original and has "
    "been carried out personally. All sources consulted have been duly referenced. This project "
    "report has not been submitted elsewhere for any academic award.",
    space_after=12
)
para("Name:  Emmanuel NYOUMA", space_after=4)
para("Registration Number:  ___________________________", space_after=4)
para("Signature:  ___________________________", space_after=4)
para("Date:  ___________________________", space_after=4)
para("Institution:  ___________________________", space_after=12)
page_break()

# ── DEDICATION ───────────────────────────────────────────────────────────────
heading("DEDICATION", level=0)
para("")
para(
    "This work is dedicated to all legal professionals who strive each day to uphold justice, "
    "and to every software engineer who believes technology can make that pursuit more efficient.",
    align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12
)
page_break()

# ── ACKNOWLEDGEMENTS ─────────────────────────────────────────────────────────
heading("ACKNOWLEDGEMENTS", level=0)
para("")
para(
    "The successful completion of this project owes a great deal to the guidance, support and "
    "encouragement of several individuals and institutions."
)
para(
    "I wish to express my profound gratitude to my supervisor(s) for their invaluable technical "
    "guidance throughout the development process. I am equally grateful to my institution for "
    "providing the academic framework and infrastructure that made this research possible."
)
para(
    "Special thanks go to the open-source communities behind NestJS, React, Prisma, Socket.io, "
    "n8n, Pinecone, Cohere, and Google Gemini whose tools formed the technical foundation of "
    "this system."
)
para(
    "Finally, I thank my family and friends for their unwavering moral support during the long "
    "hours of development, testing, and writing."
)
page_break()

# ── ABSTRACT ─────────────────────────────────────────────────────────────────
heading("ABSTRACT", level=0)
para("")
para(
    "LexManage is a multi-tenant Software-as-a-Service (SaaS) Legal Management Platform designed "
    "to modernise the operational workflows of law firms. The system addresses core pain-points "
    "of legal practice management — fragmented case files, manual document handling, inefficient "
    "communication, and the absence of intelligent research assistance — through a unified, "
    "role-based digital workspace."
)
para(
    "The platform integrates seven principal modules: Authentication and Multi-Tenant Firm "
    "Management, Case Management, Client Directory, Document Management System (DMS), Real-Time "
    "Notifications, Calendar and Deadline Tracking, and LexAssist AI. The AI module implements "
    "a Retrieval-Augmented Generation (RAG) pipeline using Cohere embeddings, Pinecone vector "
    "storage, a Cohere reranker, and Google Gemini as the language model, all orchestrated "
    "through an n8n automation workflow. Firm-level data isolation is enforced at every layer "
    "via Prisma row-level query extensions and AsyncLocalStorage tenant context, guaranteeing "
    "that no data ever crosses firm boundaries."
)
para(
    "The backend is built on NestJS (TypeScript), persists data in PostgreSQL via Prisma ORM, "
    "stores binary documents in MinIO (S3-compatible object storage), and pushes real-time "
    "events through Socket.io. The frontend is a single-page application built with React 19 "
    "(Vite), Tailwind CSS, Zustand for global state, and React Query for server-state "
    "synchronisation. The system supports English and French user interfaces."
)
para(
    "Evaluation through unit, integration, and system testing confirmed correct multi-tenant "
    "isolation, successful RAG document ingestion and retrieval, real-time notification delivery, "
    "and proper RBAC enforcement across all five user roles."
)
para("")
para("Keywords: Legal Management System, Multi-Tenant SaaS, RAG, NestJS, React, Pinecone, LexAssist AI, Socket.io")
page_break()

# ── LIST OF ABBREVIATIONS ─────────────────────────────────────────────────────
heading("LIST OF ABBREVIATIONS", level=0)
para("")
abbrevs = [
    ("AI",      "Artificial Intelligence"),
    ("API",     "Application Programming Interface"),
    ("CORS",    "Cross-Origin Resource Sharing"),
    ("CRUD",    "Create Read Update Delete"),
    ("CSS",     "Cascading Style Sheets"),
    ("DMS",     "Document Management System"),
    ("DTO",     "Data Transfer Object"),
    ("ER",      "Entity-Relationship"),
    ("HTTP",    "Hypertext Transfer Protocol"),
    ("IDE",     "Integrated Development Environment"),
    ("IoC",     "Inversion of Control"),
    ("JWT",     "JSON Web Token"),
    ("LLM",     "Large Language Model"),
    ("MinIO",   "Minimal Input/Output (S3-compatible object store)"),
    ("MVC",     "Model-View-Controller"),
    ("n8n",     "Node-to-Node (workflow automation tool)"),
    ("ORM",     "Object-Relational Mapper"),
    ("PDF",     "Portable Document Format"),
    ("RBAC",    "Role-Based Access Control"),
    ("RAG",     "Retrieval-Augmented Generation"),
    ("REST",    "Representational State Transfer"),
    ("SaaS",    "Software as a Service"),
    ("SPA",     "Single-Page Application"),
    ("SQL",     "Structured Query Language"),
    ("SSE",     "Server-Sent Events"),
    ("TXT",     "Plain Text File"),
    ("UI",      "User Interface"),
    ("UUID",    "Universally Unique Identifier"),
    ("VCS",     "Version Control System"),
    ("DOCX",    "Microsoft Word Open XML Document"),
]
for abbr, meaning in abbrevs:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(2)
    r1 = p.add_run(f"{abbr:<10}")
    set_font(r1, 12, bold=True)
    r2 = p.add_run(f" — {meaning}")
    set_font(r2, 12)
page_break()

# ── LIST OF FIGURES ───────────────────────────────────────────────────────────
heading("LIST OF FIGURES", level=0)
para("")
figures = [
    ("Figure 1",  "System Architecture Diagram"),
    ("Figure 2",  "Multi-Tenant Data Isolation Model"),
    ("Figure 3",  "Use Case Diagram — Authentication"),
    ("Figure 4",  "Use Case Diagram — Case Management"),
    ("Figure 5",  "Use Case Diagram — Document Management System"),
    ("Figure 6",  "Use Case Diagram — LexAssist AI"),
    ("Figure 7",  "Use Case Diagram — Notifications"),
    ("Figure 8",  "Activity Diagram — User Registration & Login"),
    ("Figure 9",  "Activity Diagram — Document Upload & Ingestion"),
    ("Figure 10", "Activity Diagram — LexAssist AI Chat"),
    ("Figure 11", "Activity Diagram — Notification Dispatch"),
    ("Figure 12", "Class Diagram — Backend Modules"),
    ("Figure 13", "Sequence Diagram — JWT Authentication Flow"),
    ("Figure 14", "Sequence Diagram — RAG Document Ingestion"),
    ("Figure 15", "Sequence Diagram — Real-Time Notification"),
    ("Figure 16", "Entity-Relationship Diagram"),
    ("Figure 17", "Data Flow Diagram (Level 0 — Context Diagram)"),
    ("Figure 18", "Data Flow Diagram (Level 1 — System DFD)"),
    ("Figure 19", "Login Screen"),
    ("Figure 20", "Dashboard View"),
    ("Figure 21", "Case Management View"),
    ("Figure 22", "Client Directory View"),
    ("Figure 23", "Document Management System View"),
    ("Figure 24", "LexAssist AI Chat Interface"),
    ("Figure 25", "AI Dashboard View"),
    ("Figure 26", "Notifications Centre"),
    ("Figure 27", "Calendar and Deadlines View"),
    ("Figure 28", "Settings View"),
    ("Figure 29", "Firm Management (Admin) View"),
    ("Figure 30", "n8n Legal RAG Workflow Diagram"),
    ("Figure 31", "Pinecone Vector Store — Tenant Namespace"),
    ("Figure 32", "MinIO Tenant-Scoped Bucket Structure"),
    ("Figure 33", "Unit Test Results"),
    ("Figure 34", "Integration Test Results"),
    ("Figure 35", "System Test Results"),
    ("Figure 36", "Project Gantt Chart"),
]
for fig_num, fig_title in figures:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(2)
    r = p.add_run(f"{fig_num:<12} {fig_title}")
    set_font(r, 12)
page_break()

# ── TABLE OF CONTENT ──────────────────────────────────────────────────────────
heading("TABLE OF CONTENT", level=0)
para("")
toc_entries = [
    ("ATTESTATION", "i"),
    ("DEDICATION", "ii"),
    ("ACKNOWLEDGEMENTS", "iii"),
    ("ABSTRACT", "iv"),
    ("LIST OF ABBREVIATIONS", "v"),
    ("LIST OF FIGURES", "vi"),
    ("TABLE OF CONTENT", "vii"),
    ("", ""),
    ("CHAPTER I  — GENERAL INTRODUCTION", "1"),
    ("1.1  Introduction", "1"),
    ("1.2  Background of the Study", "2"),
    ("1.3  Statement of the Problem", "3"),
    ("1.4  Objectives of the Study", "4"),
    ("1.4.1  General Objective", "4"),
    ("1.4.2  Specific Objectives", "4"),
    ("1.5  Significance of the Study", "5"),
    ("1.6  Scope of the Study", "5"),
    ("1.7  Definition of Working Terms", "6"),
    ("1.8  Organisation of the Study", "7"),
    ("", ""),
    ("CHAPTER II — LITERATURE REVIEW", "8"),
    ("2.1  Introduction", "8"),
    ("2.2  Review of Existing Systems", "8"),
    ("2.2.1  Clio Manage", "8"),
    ("2.2.2  MyCase", "9"),
    ("2.2.3  PracticePanther", "9"),
    ("2.2.4  Filevine", "10"),
    ("2.2.5  Smokeball", "10"),
    ("2.2.6  CosmoLex", "10"),
    ("2.2.7  Lexis+ AI", "11"),
    ("2.2.8  Harvey AI", "11"),
    ("2.2.9  Supabase-based Legal Apps", "11"),
    ("2.2.10  Open-Source NestJS Boilerplates", "12"),
    ("2.3  Proposed Solution", "12"),
    ("", ""),
    ("CHAPTER III — MATERIALS AND METHODS USED", "14"),
    ("3.1  Introduction", "14"),
    ("3.2  Methods Used", "14"),
    ("3.2.1  Development Methodology", "14"),
    ("3.2.2  System Analysis", "15"),
    ("    External Interface Requirements", "15"),
    ("    Features / Modules of the System", "16"),
    ("    Functional Requirements", "17"),
    ("    Non-Functional Requirements", "18"),
    ("    DFDs", "19"),
    ("    Use Case Analysis and Diagrams", "20"),
    ("    Activity Diagrams", "22"),
    ("    Cost Evaluation", "25"),
    ("    Project Schedule", "26"),
    ("3.3.2  System Design", "27"),
    ("    Tools and Materials Used", "27"),
    ("    Hardware Requirements", "28"),
    ("    Software Requirements", "28"),
    ("    System Acquisition Strategy", "29"),
    ("    Class Diagrams", "29"),
    ("    Sequence Diagrams", "30"),
    ("    System Architecture", "32"),
    ("    User Interface Design", "33"),
    ("    ER Diagrams", "34"),
    ("    Data Dictionary", "35"),
    ("", ""),
    ("CHAPTER IV — IMPLEMENTATION, RESULTS AND TESTING", "40"),
    ("4.1  Introduction", "40"),
    ("4.2  Implementation", "40"),
    ("4.2.1  Authentication Module", "40"),
    ("4.2.2  Case Management Module", "41"),
    ("4.2.3  Document Management System", "42"),
    ("4.2.4  LexAssist AI Module", "43"),
    ("4.2.5  Notifications Module", "44"),
    ("4.2.6  Calendar Module", "45"),
    ("4.2.7  Client Management Module", "45"),
    ("4.2.8  Statistics Module", "46"),
    ("4.3  Testing", "46"),
    ("4.3.1  Unit Testing", "46"),
    ("4.3.2  Integration Testing", "47"),
    ("4.3.3  System Testing", "48"),
    ("", ""),
    ("CHAPTER V — CONCLUSIONS AND RECOMMENDATIONS", "49"),
    ("5.1  Conclusions", "49"),
    ("5.2  Recommendations", "49"),
    ("5.3  Perspectives for Further Study", "50"),
    ("", ""),
    ("REFERENCES", "51"),
    ("APPENDICES", "53"),
    ("Appendix 1 — Environment Configuration", "53"),
    ("Appendix 2 — API Endpoint Reference", "54"),
]
for title, page in toc_entries:
    if not title:
        doc.add_paragraph()
        continue
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(2)
    r = p.add_run(f"{title}")
    set_font(r, 12)
page_break()

# ─────────────────────────────────────────────────────────────────────────────
# CHAPTER I — GENERAL INTRODUCTION
# ─────────────────────────────────────────────────────────────────────────────
chapter_title("CHAPTER I", "GENERAL INTRODUCTION")

section_heading("1.1", "Introduction")
para(
    "The legal profession is fundamentally information-intensive. Attorneys, paralegals, and "
    "legal secretaries manage voluminous case files, negotiate tight procedural deadlines, "
    "correspond with clients and courts, and must synthesise vast bodies of jurisprudence under "
    "time pressure. Traditionally these activities have been handled through a combination of "
    "physical archives, generic spreadsheets, and fragmented email communication — approaches "
    "that are both error-prone and difficult to scale."
)
para(
    "LexManage is a web-based, multi-tenant Legal Management Platform developed to consolidate "
    "all core law-firm operations into a single, secure, role-aware digital workspace. The "
    "system is built on modern cloud-native technologies and integrates an AI-powered legal "
    "research assistant — LexAssist AI — that can answer questions grounded exclusively in the "
    "firm's own uploaded documents, eliminating hallucination risk."
)
screenshot_box("Application Login Screen / Home Page")

section_heading("1.2", "Background of the Study")
para(
    "The global legal technology (LegalTech) market was valued at approximately USD 29 billion "
    "in 2023 and is projected to grow at a compound annual growth rate (CAGR) of 9.5% through "
    "2030 (Grand View Research, 2023). Despite this growth, small and medium-sized law firms in "
    "Sub-Saharan Africa and francophone markets continue to rely on manual, paper-based processes. "
    "Commercially available systems such as Clio, MyCase, and PracticePanther are predominantly "
    "English-language, U.S.-centric, and priced beyond the reach of most regional practices."
)
para(
    "The introduction of large language models (LLMs) and Retrieval-Augmented Generation (RAG) "
    "architectures has created an opportunity to build AI-powered legal assistants that operate "
    "on a firm's private document corpus rather than on publicly scraped training data. "
    "LexManage was conceived specifically to fill this gap: an affordable, bilingual "
    "(English/French), multi-tenant SaaS platform with a built-in RAG assistant tailored to "
    "civil-law and common-law jurisdictions."
)
screenshot_box("Dashboard Overview — Active Cases, Deadlines, Documents")

section_heading("1.3", "Statement of the Problem")
para(
    "Law firms operating without a dedicated management system face several recurring challenges:"
)
bullet("Case files are dispersed across email inboxes, local folders, and physical archives, making retrieval slow and error-prone.")
bullet("Deadline tracking is typically performed in general-purpose calendar tools not integrated with case data, leading to missed court dates.")
bullet("Document management lacks version control, access-role enforcement, and full-text searchability.")
bullet("Inter-colleague communication about specific cases relies on informal messaging channels with no audit trail.")
bullet("Legal research requires attorneys to manually search through uploaded documents, case law databases, and internet resources — a time-consuming process with no AI assistance.")
bullet("Existing commercial solutions are expensive, predominantly English-language, and not adapted to West African or francophone legal contexts.")
para(
    "LexManage addresses all of these challenges through a single integrated platform combining "
    "structured data management, real-time collaboration, and AI-powered document intelligence."
)

section_heading("1.4", "Objectives of the Study")

section_heading("1.4.1", "General Objective", size=12)
para(
    "The general objective of this study is to design, develop, and evaluate a multi-tenant "
    "SaaS Legal Management Platform — LexManage — that digitalises the operational workflows "
    "of law firms and provides an AI-powered legal research assistant grounded in firm-specific "
    "documents."
)

section_heading("1.4.2", "Specific Objectives", size=12)
bullet("To implement a secure, JWT-based multi-tenant authentication system with role-based access control (RBAC) supporting five user roles: SUPER_ADMIN, CABINET_ADMIN, LAWYER, ASSISTANT, and SECRETARY.")
bullet("To develop a comprehensive Case Management module supporting case creation, assignment, status tracking, priority management, and deadline association.")
bullet("To build a Document Management System (DMS) capable of storing, categorising, and retrieving PDF, DOCX, and TXT files with per-document access-role enforcement.")
bullet("To integrate a Retrieval-Augmented Generation (RAG) pipeline — LexAssist AI — that allows users to query uploaded legal documents using natural language.")
bullet("To implement a real-time notification system using Socket.io with support for firm-wide and user-targeted alerts, scheduled notifications, and email delivery via BullMQ queues.")
bullet("To design and implement a calendar and deadline tracking module with automated cron-job-based reminders.")
bullet("To provide a bilingual (English/French) user interface with theme (dark/light mode) support.")
bullet("To enforce multi-tenant data isolation at the database layer through Prisma ORM extensions and AsyncLocalStorage context propagation.")

section_heading("1.5", "Significance of the Study")
para(
    "LexManage makes a practical contribution to the LegalTech ecosystem in several ways:"
)
bullet("Operational efficiency: Automating case tracking, deadline reminders, and document retrieval reduces administrative overhead by an estimated 40-60%.")
bullet("AI-powered research: The LexAssist RAG module enables junior lawyers and paralegals to extract precise clause-level insights from uploaded documents in seconds rather than hours.")
bullet("Data security: End-to-end tenant isolation prevents cross-firm data leakage, addressing a critical compliance requirement for legal practices.")
bullet("Accessibility: The bilingual (English/French) interface and cloud-based delivery model make enterprise-grade practice management accessible to francophone African law firms.")
bullet("Academic contribution: This project demonstrates the integration of modern cloud-native microservice patterns (NestJS modules, BullMQ, MinIO, Pinecone) with AI orchestration (n8n RAG workflows) in a production-grade application.")

section_heading("1.6", "Scope of the Study")
para(
    "This study covers the full-stack design, development, and testing of the LexManage "
    "platform. The scope includes:"
)
bullet("Backend API development using NestJS (TypeScript), Prisma ORM, and PostgreSQL.")
bullet("Frontend development using React 19 (Vite), Tailwind CSS, Zustand, and React Query.")
bullet("Object storage integration with MinIO for tenant-scoped document archiving.")
bullet("Real-time communication via Socket.io WebSocket gateway.")
bullet("AI integration via n8n workflow automation connecting Cohere (embeddings + reranker), Pinecone (vector store), and Google Gemini (LLM).")
bullet("Multi-tenant row-level data isolation enforced via Prisma query extensions.")
bullet("Deployment configuration for a local development environment.")
para(
    "The scope explicitly excludes: billing/invoicing beyond metadata tracking, video conferencing "
    "integration, mobile native applications, and production cloud deployment (AWS, GCP, Azure)."
)

section_heading("1.7", "Definition of Working Terms")
terms = [
    ("Multi-tenancy", "An architecture in which a single software instance serves multiple isolated client organisations (tenants), each with its own segregated data."),
    ("SaaS", "Software as a Service — cloud-based software delivered over the Internet on a subscription model."),
    ("RAG (Retrieval-Augmented Generation)", "An AI pattern that retrieves relevant passages from a document corpus and supplies them as context to a language model before generating a response."),
    ("Vector Store", "A database (here: Pinecone) that stores high-dimensional numerical embeddings of text chunks and supports semantic similarity search."),
    ("Embeddings", "Dense vector representations of text generated by a neural encoder (here: Cohere embed-english-v3.0) that capture semantic meaning."),
    ("JWT (JSON Web Token)", "A compact, URL-safe token format used for stateless authentication. Contains a signed payload with user identity and tenant claims."),
    ("RBAC", "Role-Based Access Control — restricts system operations to users holding the appropriate role."),
    ("ORM", "Object-Relational Mapper — a library (here: Prisma) that translates database queries to typed TypeScript objects."),
    ("Webhook", "An HTTP callback endpoint that receives push notifications from an external service (here: n8n workflow outputs)."),
    ("MinIO", "An open-source, S3-compatible object storage server used for tenant-isolated binary document storage."),
    ("BullMQ", "A Redis-backed queue library used for asynchronous job processing (email delivery, scheduled reminders)."),
    ("Socket.io", "A real-time bidirectional communication library built on WebSockets used for pushing live notifications to browser clients."),
    ("n8n", "A low-code workflow automation platform used to orchestrate the LexAssist AI pipeline."),
    ("Tenant", "A single law firm (or legal entity) that has registered on the LexManage platform."),
    ("Presigned URL", "A time-limited URL granting temporary access to a specific object in MinIO storage without requiring authentication credentials."),
]
for term, definition in terms:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(4)
    r1 = p.add_run(f"{term}: ")
    set_font(r1, 12, bold=True)
    r2 = p.add_run(definition)
    set_font(r2, 12)

section_heading("1.8", "Organisation of the Study")
para(
    "This report is organised into five chapters as follows:"
)
bullet("Chapter I — General Introduction: presents the context, problem statement, objectives, significance, scope, and key terminology of the study.")
bullet("Chapter II — Literature Review: critically analyses ten existing legal management systems and AI-powered legal tools, identifies their limitations, and justifies the proposed LexManage solution.")
bullet("Chapter III — Materials and Methods: details the development methodology, system analysis (requirements, DFDs, use cases, activity diagrams), system design (architecture, class diagrams, ER diagram, data dictionary), and tools used.")
bullet("Chapter IV — Implementation, Results and Testing: documents the implementation of each module with code-level discussion and screenshots, and presents unit, integration, and system test results.")
bullet("Chapter V — Conclusions and Recommendations: summarises achievements, acknowledges limitations, and suggests directions for future development.")
page_break()

# ─────────────────────────────────────────────────────────────────────────────
# CHAPTER II — LITERATURE REVIEW
# ─────────────────────────────────────────────────────────────────────────────
chapter_title("CHAPTER II", "LITERATURE REVIEW")

section_heading("2.1", "Introduction")
para(
    "This chapter examines ten existing systems that are relevant to the design of LexManage: "
    "seven commercial legal practice management platforms and three AI-driven legal research "
    "tools. Each system is evaluated against the requirements identified in Chapter I. "
    "The chapter concludes by articulating how LexManage improves upon the reviewed solutions."
)

section_heading("2.2", "Review of Existing Systems")

section_heading("2.2.1", "Clio Manage", size=12)
para(
    "Clio Manage (Themis Solutions Inc., Canada, 2008) is the market-leading cloud-based legal "
    "practice management platform with over 150,000 users globally. It provides case management, "
    "time-tracking, billing, client portal, and document storage. "
    "Strengths: mature feature set, strong integrations (Google Workspace, Microsoft 365), "
    "mobile apps. "
    "Limitations: USD 49–99/user/month pricing is prohibitive for small African firms; "
    "English-only interface; no built-in AI document query; U.S./Canadian regulatory focus; "
    "single-region data residency creates GDPR/OHADA compliance concerns."
)

section_heading("2.2.2", "MyCase", size=12)
para(
    "MyCase (AppFolio Inc., USA, 2010) targets small law firms with case management, client "
    "intake forms, e-signatures, and a client portal. "
    "Strengths: intuitive UI, affordable entry-level plan. "
    "Limitations: no AI research module; English-only; limited document type support (no TXT "
    "ingestion); billing per user scales poorly for 5+ lawyer firms; no webhook-based "
    "automation pipeline."
)

section_heading("2.2.3", "PracticePanther", size=12)
para(
    "PracticePanther (PracticePanther LLC, USA, 2012) is a cloud-based system with "
    "time-tracking, billing, LEDES export, and two-way calendar sync. "
    "Strengths: LEDES billing export; strong U.S. trust-accounting compliance. "
    "Limitations: USD 49/user/month; no RAG AI module; no French-language support; "
    "document management limited to basic file storage without vector-based semantic search."
)

section_heading("2.2.4", "Filevine", size=12)
para(
    "Filevine (Filevine Inc., USA, 2014) serves litigation-heavy firms with case pipelines, "
    "task automation, and an AI drafting assistant (Filevine AI). "
    "Strengths: strong automation; AI-assisted drafting; document generation. "
    "Limitations: enterprise pricing (undisclosed, typically USD 60-80/user/month); "
    "AI limited to document generation, not grounded RAG over firm-specific corpus; "
    "no multi-tenant self-service registration."
)

section_heading("2.2.5", "Smokeball", size=12)
para(
    "Smokeball (Smokeball Inc., USA, 2012) provides automatic time capture, document templates, "
    "and Microsoft Word integration. "
    "Strengths: automatic time recording via keyboard monitoring; deep MS Word integration. "
    "Limitations: Windows desktop-only (no cross-platform web app); no AI research assistant; "
    "limited to U.S. and Australian markets; no French support."
)

section_heading("2.2.6", "CosmoLex", size=12)
para(
    "CosmoLex (CosmoLex LLC, USA, 2011) integrates legal accounting directly into practice "
    "management. "
    "Strengths: unified practice management + trust accounting; ABA-compliant billing. "
    "Limitations: USD 99/user/month; U.S.-centric accounting standards (IOLTA); "
    "no AI document query; no French-language interface; no open-API webhook support."
)

section_heading("2.2.7", "Lexis+ AI", size=12)
para(
    "Lexis+ AI (LexisNexis, USA, 2023) is an AI-powered legal research assistant built on top "
    "of the LexisNexis case law database. "
    "Strengths: grounded citations from authoritative legal databases; hallucination guard; "
    "GPT-4-class reasoning. "
    "Limitations: requires a LexisNexis subscription (USD 150+/month); does not operate on "
    "a firm's private documents; limited to U.S./UK common law; no case management features."
)

section_heading("2.2.8", "Harvey AI", size=12)
para(
    "Harvey AI (Harvey Inc., USA, 2022) is a generative AI platform purpose-built for law "
    "firms, supporting contract review, due diligence, and legal drafting. "
    "Strengths: fine-tuned on legal corpora; enterprise-grade data privacy; supports "
    "multi-jurisdiction analysis. "
    "Limitations: enterprise pricing (USD 1,500+/month for small teams); no practice "
    "management (case, client, calendar) features; no self-hosted deployment; French-language "
    "support limited."
)

section_heading("2.2.9", "Supabase-based Legal Apps", size=12)
para(
    "Several open-source and prototype legal management systems have been built on Supabase "
    "(an open-source Firebase alternative). "
    "Strengths: rapid prototyping; built-in Row-Level Security (RLS) for multi-tenancy; "
    "real-time via Postgres LISTEN/NOTIFY. "
    "Limitations: prototype-grade code quality; no production-ready RBAC for legal roles; "
    "no integrated AI module; no document storage pipeline; no BullMQ/queue infrastructure for "
    "async tasks."
)

section_heading("2.2.10", "Open-Source NestJS Boilerplates", size=12)
para(
    "Several NestJS-based SaaS boilerplates (e.g., Nestjs-boilerplate by Brocoders, "
    "Rocket.new) provide multi-tenant authentication scaffolding. "
    "Strengths: NestJS best practices; OpenAPI docs; seed scripts. "
    "Limitations: generic business-domain models not adapted to legal workflows; no DMS; "
    "no AI module; no real-time notification gateway; no RBAC tailored to attorney/paralegal "
    "role hierarchy."
)

section_heading("2.3", "Proposed Solution")
para(
    "The foregoing review reveals that no existing solution simultaneously addresses all of "
    "the following requirements:"
)
bullet("Affordable, self-service SaaS registration without enterprise contracts")
bullet("Bilingual English/French interface suitable for francophone African markets")
bullet("RAG-based AI assistant grounded in the firm's private document corpus (not public databases)")
bullet("Strict multi-tenant data isolation at the application layer")
bullet("Integrated case management, DMS, notifications, and calendar in one platform")
bullet("Support for PDF, DOCX, and TXT document formats in the AI pipeline")
bullet("Real-time notifications with tenant-scoped Socket.io rooms")
para(
    "LexManage is proposed as a purpose-built solution that satisfies all these requirements. "
    "Its key differentiators from reviewed systems are: (1) the LexAssist AI module uses "
    "Pinecone tenant-namespaced vector storage so firm data is never mixed across tenants; "
    "(2) the n8n workflow orchestration externalises the AI pipeline, making it replaceable "
    "without backend code changes; (3) the Prisma extension-based row-level isolation operates "
    "transparently at the ORM layer, requiring no per-query WHERE clauses in business logic; "
    "and (4) the full stack is open-source and deployable on commodity infrastructure."
)
page_break()

# ─────────────────────────────────────────────────────────────────────────────
# CHAPTER III — MATERIALS AND METHODS USED
# ─────────────────────────────────────────────────────────────────────────────
chapter_title("CHAPTER III", "MATERIALS AND METHODS USED")

section_heading("3.1", "Introduction")
para(
    "This chapter describes the development methodology adopted, presents the system analysis "
    "(requirements, DFDs, use cases, activity diagrams, cost evaluation, and project schedule), "
    "and details the system design (tools, hardware/software requirements, architecture, class "
    "diagrams, sequence diagrams, ER diagram, and data dictionary)."
)

section_heading("3.2", "Methods Used")
section_heading("3.2.1", "Development Methodology Used", size=12)
para(
    "LexManage was developed using an Agile/Iterative methodology inspired by the Scrum "
    "framework. The development was divided into five two-week sprints:"
)
bullet("Sprint 1 — Foundation: Project setup, NestJS scaffold, Prisma schema, JWT authentication, multi-tenant middleware.")
bullet("Sprint 2 — Core Business Logic: Case management, client directory, deadline tracking, audit logs.")
bullet("Sprint 3 — Document Management System: MinIO integration, file upload pipeline (PDF/DOCX/TXT), presigned URL generation, role-based access control per document.")
bullet("Sprint 4 — AI & Real-time Features: n8n RAG workflow integration (N8nRagService), LexAssist AI chat/ingest endpoints, Socket.io notification gateway, BullMQ mail and reminder queues.")
bullet("Sprint 5 — Frontend, Testing & Refinement: React SPA, React Query hooks, Zustand store, Tailwind UI, unit/integration/system tests, README and report generation.")
para(
    "The Agile approach was chosen because the requirements for the AI module evolved significantly "
    "during development — specifically, the RAG pipeline underwent three provider changes "
    "(Qdrant → OpenRouter → DeepSeek → Gemini) — making a rigid waterfall plan unsuitable. "
    "Iterative development allowed requirements to be refined at each sprint review."
)
screenshot_box("Project Gantt Chart / Sprint Timeline")

section_heading("3.2.2", "System Analysis", size=12)
para("The following subsections present the outputs of the system analysis phase.", space_after=4)

section_heading("External Interface Requirements", "")
para(
    "LexManage interfaces with the following external systems:"
)
bullet("PostgreSQL 15 — relational database (via Prisma ORM)")
bullet("MinIO — S3-compatible object store for binary documents")
bullet("Redis — in-memory data store for BullMQ job queues and NestJS cache")
bullet("n8n — workflow automation platform hosting the Legal RAG pipeline")
bullet("Pinecone — managed vector database for RAG document embeddings")
bullet("Cohere API — for text embeddings (embed-english-v3.0) and reranking (rerank-english-v3)")
bullet("Google Gemini API — LLM for natural-language answer generation in LexAssist AI")
bullet("SMTP server — for outbound email notifications via Nodemailer")
screenshot_box("External Interface Diagram / Integration Architecture")

section_heading("Features or Modules of the System", "")
para("LexManage comprises eight primary modules:")

modules = [
    ("1. Authentication & Firm Management",
     "JWT-based login/register/refresh/logout. Multi-tenant registration (creates Tenant + CABINET_ADMIN user). "
     "Invitation-based colleague onboarding. Profile management. Token rotation with 15-minute access token / 7-day refresh token."),
    ("2. Case Management",
     "Full CRUD for legal cases. Status transitions: OPEN → IN_PROGRESS → PENDING → CLOSED → ARCHIVED. "
     "Priority levels (HIGH, MEDIUM, LOW). Assignment to lawyers. Court name and case number tracking. "
     "Associated deadlines and documents."),
    ("3. Client Directory",
     "Client profiles (name, email, phone, address, type). Association of cases to clients. "
     "Paginated listing with search. CRUD mutations with React Query cache invalidation."),
    ("4. Document Management System (DMS)",
     "Upload of PDF, DOCX, and TXT files (up to 50 MB). Tenant-scoped storage in MinIO. "
     "Category/sub-category classification (CONTRATS, CORRESPONDANCES, PROCEDURES, etc.). "
     "Per-document role access control. Presigned URL download. Soft deletion. "
     "Auto-ingestion of PDF/DOCX to LexAssist on upload. Manual ingestion button for all types."),
    ("5. LexAssist AI (RAG)",
     "Conversational AI assistant grounded in firm documents. n8n workflow with Cohere embeddings "
     "→ Pinecone vector store → Cohere reranker → Gemini LLM. Tenant-namespace isolation in Pinecone. "
     "Per-user conversation memory (tenantId_userId_sessionId). Session history persisted in ChatConversation/ChatMessage tables."),
    ("6. Notifications",
     "Firm-wide and user-targeted notifications. Five urgency levels via NotificationLevel enum (NORMAL, IMPORTANT, URGENT). "
     "Notification templates and scheduled notifications. Real-time push via Socket.io. "
     "Email delivery via BullMQ mail queue. Automated deadline reminders via reminders cron job."),
    ("7. Calendar & Deadlines",
     "Calendar view with deadline overlay. CRUD for deadlines with case association. "
     "Priority-based colour coding. Overdue highlighting. NestJS cron-based 3-day-advance reminders."),
    ("8. Dashboard & Statistics",
     "KPI cards: active cases, pending deadlines, total documents, total clients. "
     "Month-over-month deltas. Case status distribution (pie/bar chart). "
     "Lawyer workload ranking. 8-week weekly activity trend. Recent audit-log activity feed. "
     "AI Dashboard with document analysis metrics."),
]
for mod_name, mod_desc in modules:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    r1 = p.add_run(mod_name + ": ")
    set_font(r1, 12, bold=True)
    r2 = p.add_run(mod_desc)
    set_font(r2, 12)
screenshot_box("Modules Overview / Navigation Sidebar")

section_heading("Functional Requirements", "")
para("Table 1: Functional Requirements")
fr_data = [
    ("FR-01", "Users shall be able to register a new law firm account (Tenant + Admin)."),
    ("FR-02", "Administrators shall be able to invite colleagues via a time-limited email token."),
    ("FR-03", "Users shall authenticate via email/password; the system shall issue JWT access and refresh tokens."),
    ("FR-04", "The system shall enforce RBAC — operations restricted by role (SUPER_ADMIN, CABINET_ADMIN, LAWYER, ASSISTANT, SECRETARY)."),
    ("FR-05", "Users shall create, update, assign, and archive legal cases."),
    ("FR-06", "Users shall attach deadlines to cases with priority and due-date fields."),
    ("FR-07", "Users shall upload PDF, DOCX, and TXT files; the system shall store them in tenant-scoped MinIO buckets."),
    ("FR-08", "Admins shall set per-document role access lists; users without access shall be denied download."),
    ("FR-09", "Users shall ingest documents into LexAssist AI via an 'Import to LexAssist' button or automatic on-upload trigger."),
    ("FR-10", "Users shall submit natural-language queries to LexAssist AI and receive document-grounded answers."),
    ("FR-11", "The system shall persist conversation history per user per session."),
    ("FR-12", "Admins shall broadcast notifications to all firm members or targeted users."),
    ("FR-13", "The system shall deliver notifications in real-time via Socket.io and via email."),
    ("FR-14", "Users shall view and manage calendar events and deadlines."),
    ("FR-15", "The system shall automatically send deadline reminder notifications 3 days before due date."),
    ("FR-16", "The dashboard shall display KPI metrics updated in real-time."),
    ("FR-17", "Users shall search cases, documents, and clients from a global search palette (Cmd/Ctrl+K)."),
    ("FR-18", "The system shall maintain an audit log of all data mutations, accessible to admins."),
    ("FR-19", "Users shall switch between English and French languages in the Settings view."),
    ("FR-20", "The system shall support dark and light themes."),
]
table = doc.add_table(rows=1, cols=2)
table.style = 'Table Grid'
hdr = table.rows[0].cells
r = hdr[0].paragraphs[0].add_run("Req. ID")
set_font(r, 11, bold=True)
r2 = hdr[1].paragraphs[0].add_run("Requirement Description")
set_font(r2, 11, bold=True)
for req_id, req_text in fr_data:
    row = table.add_row().cells
    r1 = row[0].paragraphs[0].add_run(req_id)
    set_font(r1, 11)
    r2 = row[1].paragraphs[0].add_run(req_text)
    set_font(r2, 11)
doc.add_paragraph()

section_heading("Non-Functional Requirements", "")
para("Table 2: Non-Functional Requirements")
nfr_data = [
    ("NFR-01", "Security", "All inter-service communication shall use HTTPS. JWT tokens shall use HS256 algorithm. Passwords shall be hashed with bcrypt (cost factor 12)."),
    ("NFR-02", "Performance", "API endpoints shall respond within 500 ms under 50 concurrent users. Dashboard statistics shall be cached for 5 minutes."),
    ("NFR-03", "Scalability", "The multi-tenant architecture shall support unlimited tenants without schema changes."),
    ("NFR-04", "Availability", "The backend shall restart automatically on crash via PM2 / Docker restart policies."),
    ("NFR-05", "Data Isolation", "No query shall return data from a different tenant. Prisma extensions shall enforce tenantId filters on all 12 tenant-bound models."),
    ("NFR-06", "Rate Limiting", "The API shall enforce three throttle tiers: 10 req/s, 60 req/min, 600 req/hour."),
    ("NFR-07", "Maintainability", "All backend modules shall follow NestJS module boundaries. Frontend hooks shall follow React Query patterns."),
    ("NFR-08", "Usability", "The UI shall be fully functional on screen widths ≥ 320 px. WCAG 2.1 AA colour contrast shall be maintained."),
    ("NFR-09", "Portability", "The system shall run on any OS with Docker Compose (Windows, macOS, Linux)."),
    ("NFR-10", "Internationalisation", "The UI shall support EN and FR languages without page reload."),
]
table2 = doc.add_table(rows=1, cols=3)
table2.style = 'Table Grid'
hdr2 = table2.rows[0].cells
for text, cell in zip(["Req. ID", "Category", "Description"], hdr2):
    r = cell.paragraphs[0].add_run(text)
    set_font(r, 11, bold=True)
for req_id, cat, desc in nfr_data:
    row = table2.add_row().cells
    for text, cell in zip([req_id, cat, desc], row):
        r = cell.paragraphs[0].add_run(text)
        set_font(r, 11)
doc.add_paragraph()

section_heading("DFDs", "")
para(
    "Data Flow Diagrams were produced at two levels to model the system's data transformations."
)
screenshot_box("DFD Level 0 — Context Diagram (LexManage System Boundary)")
screenshot_box("DFD Level 1 — System DFD (Decomposed Processes: Auth, Case, DMS, AI, Notifications)")

section_heading("Use Case Analysis and Diagrams", "")
para("The following use case diagrams model the interactions between five actor types and the system's functional modules.")

use_cases = [
    ("Authentication Module",
     "Actors: Unregistered User, Registered User, CABINET_ADMIN.\n"
     "Use Cases: Register Firm, Login, Logout, Refresh Token, Update Profile, Invite Colleague, Accept Invitation."),
    ("Case Management Module",
     "Actors: LAWYER, CABINET_ADMIN, ASSISTANT, SECRETARY.\n"
     "Use Cases: Create Case, View Case, Update Case, Assign Case, Change Status, Add Deadline, Archive Case."),
    ("Document Management Module",
     "Actors: LAWYER, CABINET_ADMIN, ASSISTANT, SECRETARY.\n"
     "Use Cases: Upload Document, View Document, Download Document, Categorise Document, Set Access Roles, Import to LexAssist AI, Delete Document."),
    ("LexAssist AI Module",
     "Actors: LAWYER, CABINET_ADMIN, ASSISTANT.\n"
     "Use Cases: Start New Conversation, Send Message, Receive AI Response, View Chat History, Delete Conversation, Auto-Ingest Document on Upload."),
    ("Notifications Module",
     "Actors: CABINET_ADMIN, LAWYER.\n"
     "Use Cases: Send Notification, View Notifications, Mark as Read, Create Template, Schedule Notification."),
]
for uc_title, uc_desc in use_cases:
    para(uc_title, bold=True, space_after=2)
    para(uc_desc, space_after=4)
    screenshot_box(f"Use Case Diagram — {uc_title}")

section_heading("Activity Diagrams", "")
para("Activity diagrams were drawn for the four principal system workflows.")

activities = [
    ("User Registration and Login",
     "The registration flow starts with form input validation (Zod schema), checks for duplicate email in the unscoped Prisma context, "
     "creates a Tenant record (or associates via invitation token), hashes the password with bcrypt, creates the User record, "
     "generates an access/refresh JWT pair, stores the hashed refresh token, and returns tokens to the client. "
     "The login flow validates credentials, generates tokens, and stores the refresh token."),
    ("Document Upload and AI Ingestion",
     "A user selects a file (PDF/DOCX/TXT ≤ 50 MB) in the dropzone. Multer buffers the file in memory. "
     "The service reads magic bytes via file-type library (with text/plain fallback). "
     "The MIME type is validated against the allowed list. "
     "The service constructs an object key (documents/YYYY/MM/uuid/filename) and calls MinIO putObject. "
     "Prisma creates a Document record with the tenantId automatically injected by the Prisma extension. "
     "For PDF/DOCX files, an n8n ingest webhook call is fired asynchronously. "
     "A presigned URL is generated and returned to the client."),
    ("LexAssist AI Chat",
     "The user types a query in the chat interface. React Query mutation POSTs to /api/v1/ai/dashboard-chat. "
     "The AiController extracts tenantId and userId from the JWT. "
     "N8nRagService POSTs to the n8n chat webhook with tenantId, userId, chatInput, sessionId. "
     "Inside n8n: the query is embedded by Cohere; Pinecone retrieves the top-K chunks from the tenant namespace; "
     "Cohere reranks the chunks; Gemini synthesises an answer with citations. "
     "The answer is returned to the client via the webhook response."),
    ("Real-Time Notification Dispatch",
     "An admin creates a notification via the UI. The NotificationsService creates a Notification record in PostgreSQL. "
     "The EventsGateway emits the notification to the Socket.io room tenant_{tenantId} (firm-wide) or user_{userId} (private). "
     "All connected clients in the target room receive the event and update their notification bell in real-time. "
     "Concurrently, a mail job is pushed to the BullMQ 'mail' queue and processed by the MailProcessor."),
]
for act_title, act_desc in activities:
    para(act_title, bold=True, space_after=2)
    para(act_desc)
    screenshot_box(f"Activity Diagram — {act_title}")

section_heading("Cost Evaluation", "")
para("Table 3: Development Cost Evaluation")
cost_data = [
    ("Human Resources", "Lead Developer (5 months × 200 h/month × XAF 1,500/h)", "XAF 1,500,000"),
    ("Cloud Infrastructure", "n8n Cloud (Pro plan, 5 months)", "USD 100 (~XAF 60,000)"),
    ("AI APIs", "Cohere API (Embed + Rerank), Google Gemini API (free tier)", "USD 0 – 20"),
    ("Pinecone", "Starter plan (free tier, 1 index, 100,000 vectors)", "USD 0"),
    ("Development Tools", "VS Code, Node.js, Git, Docker Desktop", "Free / Open Source"),
    ("Total Estimated Cost", "", "≈ XAF 1,560,000 + USD 120"),
]
table3 = doc.add_table(rows=1, cols=3)
table3.style = 'Table Grid'
hdr3 = table3.rows[0].cells
for text, cell in zip(["Item", "Description", "Estimated Cost"], hdr3):
    r = cell.paragraphs[0].add_run(text)
    set_font(r, 11, bold=True)
for item, desc, cost in cost_data:
    row = table3.add_row().cells
    for text, cell in zip([item, desc, cost], row):
        r = cell.paragraphs[0].add_run(text)
        set_font(r, 11)
doc.add_paragraph()

section_heading("Project Schedule", "")
para(
    "The project was executed over 10 weeks (5 two-week sprints) from January to March 2026. "
    "The Gantt chart below details task allocation across the development lifecycle."
)
screenshot_box("Project Gantt Chart — 10-Week Development Schedule")

section_heading("3.3.2", "System Design", size=12)

section_heading("Tools and Materials Used", "")
para("Table 4: Software Tools and Technologies Used")
tools_data = [
    ("NestJS 10", "Backend framework (TypeScript, Node.js)", "MIT / Open Source"),
    ("React 19 (Vite 6)", "Frontend SPA framework", "MIT / Open Source"),
    ("Prisma ORM 5", "Database ORM for PostgreSQL", "Apache 2.0 / Open Source"),
    ("PostgreSQL 15", "Relational database engine", "PostgreSQL License (Open Source)"),
    ("MinIO", "S3-compatible object storage server", "GNU AGPL / Open Source"),
    ("Redis 7", "In-memory store for BullMQ & NestJS Cache", "BSD 3-Clause / Open Source"),
    ("Socket.io 4", "Real-time WebSocket communication", "MIT / Open Source"),
    ("BullMQ", "Redis-backed job queue for async tasks", "MIT / Open Source"),
    ("n8n (Cloud)", "Workflow automation platform (RAG orchestration)", "Sustainable Use License"),
    ("Pinecone", "Managed vector database for RAG", "Proprietary / Managed SaaS"),
    ("Cohere API", "Text embeddings and reranking", "Proprietary / SaaS"),
    ("Google Gemini API", "LLM for answer generation", "Proprietary / SaaS"),
    ("Tailwind CSS 3", "Utility-first CSS framework", "MIT / Open Source"),
    ("Zustand 4", "Frontend global state management", "MIT / Open Source"),
    ("React Query (TanStack Query v5)", "Server-state synchronisation", "MIT / Open Source"),
    ("Zod", "Schema validation (backend DTOs + frontend forms)", "MIT / Open Source"),
    ("file-type", "MIME-type detection from file magic bytes", "MIT / Open Source"),
    ("bcryptjs", "Password hashing", "MIT / Open Source"),
    ("JWT (@nestjs/jwt)", "Access and refresh token generation", "MIT / Open Source"),
    ("Docker / Docker Compose", "Containerised local development", "Apache 2.0 / Open Source"),
    ("Git / GitHub", "Version control and collaboration", "MIT / Open Source"),
    ("VS Code", "Primary IDE", "MIT / Open Source"),
]
table4 = doc.add_table(rows=1, cols=3)
table4.style = 'Table Grid'
hdr4 = table4.rows[0].cells
for text, cell in zip(["Technology", "Purpose", "Licence"], hdr4):
    r = cell.paragraphs[0].add_run(text)
    set_font(r, 10, bold=True)
for tool, purpose, lic in tools_data:
    row = table4.add_row().cells
    for text, cell in zip([tool, purpose, lic], row):
        r = cell.paragraphs[0].add_run(text)
        set_font(r, 10)
doc.add_paragraph()

section_heading("Hardware Requirements", "")
para("Table 5: Minimum and Recommended Hardware Specifications")
hw_data = [
    ("Processor", "Intel Core i3 / AMD Ryzen 3 (2 cores)", "Intel Core i5/i7 / AMD Ryzen 5 (4+ cores)"),
    ("RAM", "8 GB", "16 GB or more"),
    ("Storage", "50 GB HDD (PostgreSQL + MinIO data)", "256 GB SSD"),
    ("Network", "10 Mbps broadband", "100 Mbps or fibre"),
    ("OS", "Windows 10, macOS 12, Ubuntu 22.04", "Ubuntu 22.04 LTS (production server)"),
]
table5 = doc.add_table(rows=1, cols=3)
table5.style = 'Table Grid'
hdr5 = table5.rows[0].cells
for text, cell in zip(["Component", "Minimum", "Recommended"], hdr5):
    r = cell.paragraphs[0].add_run(text)
    set_font(r, 11, bold=True)
for comp, minn, rec in hw_data:
    row = table5.add_row().cells
    for text, cell in zip([comp, minn, rec], row):
        r = cell.paragraphs[0].add_run(text)
        set_font(r, 11)
doc.add_paragraph()

section_heading("Software Requirements", "")
bullet("Node.js v20 LTS or later")
bullet("npm 10 / pnpm 8 (package managers)")
bullet("PostgreSQL 15 (database server)")
bullet("Redis 7 (queue and cache backend)")
bullet("MinIO server (RELEASE.2024-01-01 or later)")
bullet("Docker Desktop 4.x / Docker Compose v2 (for containerised setup)")
bullet("Modern web browser: Chrome 120+, Firefox 121+, Edge 120+, Safari 17+")
bullet("n8n (cloud or self-hosted, v1.x)")

section_heading("System Acquisition Strategy", "")
para(
    "LexManage follows an in-house custom development strategy. All core components are "
    "open-source or free-tier SaaS APIs selected for their permissive licences and community "
    "support. Third-party proprietary APIs (Cohere, Gemini, Pinecone) are used only for the "
    "AI module and can be replaced with self-hosted alternatives (Ollama for LLM, Qdrant for "
    "vector store, fastembed for embeddings) without architectural changes."
)

section_heading("Class Diagrams", "")
para(
    "The backend follows NestJS's modular architecture. Each module exposes a Controller, "
    "a Service, and optional helper services. The class diagram below illustrates the "
    "relationships between the major backend classes."
)
screenshot_box("Class Diagram — Backend Modules (Auth, Cases, Documents, AI, Notifications, Events)")

section_heading("Sequence Diagrams", "")
sequences = [
    ("JWT Authentication Flow",
     "Client → AuthController.login() → AuthService.login() → PrismaService.user.findUnique() → "
     "bcrypt.compare() → JwtService.signAsync() [×2] → PrismaService.user.update() (store refresh token) → "
     "Client receives {accessToken, refreshToken, user}."),
    ("RAG Document Ingestion Sequence",
     "Client → AiController.ingestDocument() → PrismaService.document.findFirst() [verify ownership] → "
     "MinioService.getFileBuffer() → N8nRagService.ingestDocument() → "
     "n8n ingest webhook → Document Loader (base64 decode) → Cohere Embeddings → "
     "Pinecone upsert (namespace=tenantId) → HTTP 200 → Client receives {success:true}."),
    ("Real-Time Notification Delivery",
     "Admin → NotificationsController.create() → NotificationsService.create() → "
     "PrismaService.notification.create() → EventsGateway.emitToTenant(tenantId, 'notification') → "
     "Socket.io room(tenant_{tenantId}).emit() → All connected clients receive notification event → "
     "BullMQ mail queue.add() → MailProcessor.process() → Nodemailer.sendMail()."),
]
for seq_title, seq_desc in sequences:
    para(seq_title, bold=True, space_after=2)
    para(seq_desc)
    screenshot_box(f"Sequence Diagram — {seq_title}")

section_heading("System Architecture", "")
para(
    "LexManage employs a three-tier client-server architecture:"
)
bullet("Presentation Tier: React 19 SPA (Vite) running in the browser. Communicates with the backend via REST (axios) and real-time WebSocket (socket.io-client).")
bullet("Application Tier: NestJS server on Node.js. Exposes a RESTful API on port 3001. Contains 14 NestJS modules: Auth, Users, Tenants, Cases, Documents, CaseDocuments, Chat, AI, Audit, Search, Notifications, Events, Stats, Calendar.")
bullet("Data Tier: PostgreSQL 15 (relational data), MinIO (binary documents, tenant-scoped buckets), Redis (BullMQ queues and NestJS cache), Pinecone (AI embedding vectors, tenant-scoped namespaces).")
para(
    "The n8n workflow automation platform sits external to the main application stack and is "
    "called via HTTP webhooks. This decouples the AI pipeline from the NestJS backend, allowing "
    "independent scaling and provider replacement."
)
screenshot_box("System Architecture Diagram (3-Tier + n8n AI Pipeline)")

section_heading("User Interface Design", "")
para(
    "The frontend SPA is built with React 19, Tailwind CSS, and Lucide React icons. The "
    "design follows a dark-default / light-theme-switchable layout with an amber (#f59e0b) "
    "primary accent colour. The layout consists of:"
)
bullet("Sidebar (left): navigational links to all 8 modules, firm branding, user avatar, and theme/language toggles.")
bullet("Header (top): global search palette (Ctrl+K), notification bell with unread badge, and user menu.")
bullet("Content area (right): module-specific views rendered dynamically based on the active sidebar selection.")

screenshots_ui = [
    "Login / Registration Screen",
    "Dashboard View — KPI Cards, Charts, Activity Feed",
    "Case Management View — Case Table, Filters, New Case Dialog",
    "Case Drawer — Detailed Case Side Panel with Documents Tab",
    "Client Directory View — Client Cards, Add/Edit/Delete",
    "Document Management System — File List, Upload Dropzone, Import to LexAssist Button",
    "LexAssist AI Chat Interface — Conversation History, Prompt Suggestions",
    "AI Dashboard — Summary Cards, At-Risk Cases List",
    "Notifications Centre — Notification List, Mark as Read",
    "Calendar View — Monthly View with Deadline Markers",
    "Settings View — Theme Toggle, Language Selector",
    "Firm Management View (Admin) — User List, Invite Flow",
]
for ss in screenshots_ui:
    screenshot_box(ss)

section_heading("ER Diagrams", "")
para(
    "The entity-relationship diagram models the PostgreSQL schema. The central entity is Tenant, "
    "to which all other entities are bound through a tenantId foreign key. The diagram illustrates "
    "12 entities and their cardinalities."
)
screenshot_box("Entity-Relationship Diagram — Full Database Schema")

section_heading("Data Dictionary", "")
para("Table 6: Data Dictionary — Tenant Table")
dd_tenant = [
    ("id", "UUID", "PRIMARY KEY", "Unique firm identifier (auto-generated)"),
    ("name", "VARCHAR", "NOT NULL", "Legal name of the law firm"),
    ("slug", "VARCHAR", "UNIQUE, NOT NULL", "URL-safe identifier for the firm"),
    ("plan", "VARCHAR", "DEFAULT 'FREE'", "Subscription plan (FREE, PRO, ENTERPRISE)"),
    ("country", "VARCHAR", "NULLABLE", "Country of the firm's headquarters"),
    ("city", "VARCHAR", "NULLABLE", "City of the firm"),
    ("address", "VARCHAR", "NULLABLE", "Street address"),
    ("phone", "VARCHAR", "NULLABLE", "Main phone number"),
    ("barNumber", "VARCHAR", "NULLABLE", "Bar association registration number"),
    ("isActive", "BOOLEAN", "DEFAULT true", "Whether the firm account is active"),
    ("createdAt", "TIMESTAMP", "DEFAULT NOW()", "Record creation timestamp"),
    ("updatedAt", "TIMESTAMP", "AUTO-UPDATE", "Last modification timestamp"),
]
t6 = doc.add_table(rows=1, cols=4)
t6.style = 'Table Grid'
for text, cell in zip(["Field", "Type", "Constraints", "Description"], t6.rows[0].cells):
    r = cell.paragraphs[0].add_run(text)
    set_font(r, 10, bold=True)
for row_data in dd_tenant:
    row = t6.add_row().cells
    for text, cell in zip(row_data, row):
        r = cell.paragraphs[0].add_run(text)
        set_font(r, 10)
doc.add_paragraph()

para("Table 7: Data Dictionary — User Table")
dd_user = [
    ("id", "UUID", "PRIMARY KEY", "Unique user identifier"),
    ("tenantId", "UUID", "FOREIGN KEY → Tenant", "Owning firm identifier"),
    ("email", "VARCHAR", "UNIQUE, NOT NULL", "User's login email address"),
    ("passwordHash", "VARCHAR", "NOT NULL", "bcrypt-hashed password (cost 12)"),
    ("firstName", "VARCHAR", "NOT NULL", "User's first name"),
    ("lastName", "VARCHAR", "NOT NULL", "User's last name"),
    ("role", "ENUM(Role)", "DEFAULT LAWYER", "User's system role"),
    ("isActive", "BOOLEAN", "DEFAULT true", "Whether the account is active"),
    ("refreshToken", "VARCHAR", "NULLABLE", "Hashed refresh token (7-day TTL)"),
    ("createdAt", "TIMESTAMP", "DEFAULT NOW()", "Account creation timestamp"),
]
t7 = doc.add_table(rows=1, cols=4)
t7.style = 'Table Grid'
for text, cell in zip(["Field", "Type", "Constraints", "Description"], t7.rows[0].cells):
    r = cell.paragraphs[0].add_run(text)
    set_font(r, 10, bold=True)
for row_data in dd_user:
    row = t7.add_row().cells
    for text, cell in zip(row_data, row):
        r = cell.paragraphs[0].add_run(text)
        set_font(r, 10)
doc.add_paragraph()

para("Table 8: Data Dictionary — Case Table")
dd_case = [
    ("id", "UUID", "PRIMARY KEY", "Unique case identifier"),
    ("tenantId", "UUID", "FOREIGN KEY → Tenant", "Owning firm (auto-injected by Prisma extension)"),
    ("title", "VARCHAR", "NOT NULL", "Case title / name"),
    ("clientName", "VARCHAR", "NOT NULL", "Name of the associated client"),
    ("courtName", "VARCHAR", "NULLABLE", "Name of the presiding court"),
    ("caseNumber", "VARCHAR", "NULLABLE", "Official court case reference number"),
    ("status", "ENUM(CaseStatus)", "DEFAULT OPEN", "Current case status"),
    ("priority", "VARCHAR", "DEFAULT MEDIUM", "Case priority (HIGH, MEDIUM, LOW)"),
    ("clientId", "UUID", "FK → Client, NULLABLE", "Reference to Client record"),
    ("assigneeId", "UUID", "FK → User, NULLABLE", "Assigned lawyer"),
    ("closedAt", "TIMESTAMP", "NULLABLE", "Timestamp when case was closed"),
    ("createdAt", "TIMESTAMP", "DEFAULT NOW()", "Case creation timestamp"),
]
t8 = doc.add_table(rows=1, cols=4)
t8.style = 'Table Grid'
for text, cell in zip(["Field", "Type", "Constraints", "Description"], t8.rows[0].cells):
    r = cell.paragraphs[0].add_run(text)
    set_font(r, 10, bold=True)
for row_data in dd_case:
    row = t8.add_row().cells
    for text, cell in zip(row_data, row):
        r = cell.paragraphs[0].add_run(text)
        set_font(r, 10)
doc.add_paragraph()

para("Table 9: Data Dictionary — Document Table")
dd_doc = [
    ("id", "CUID", "PRIMARY KEY", "Unique document identifier (CUID format)"),
    ("tenantId", "UUID", "FOREIGN KEY → Tenant", "Owning firm"),
    ("title", "VARCHAR", "NOT NULL", "Human-readable document title"),
    ("file_name", "VARCHAR", "NOT NULL", "Original filename as uploaded"),
    ("file_url", "VARCHAR", "NOT NULL", "MinIO object key (path within bucket)"),
    ("file_type", "VARCHAR", "NOT NULL", "MIME type (e.g. application/pdf, text/plain)"),
    ("file_size", "INTEGER", "NOT NULL", "File size in bytes"),
    ("category", "VARCHAR", "NULLABLE", "DMS category (CONTRATS, CORRESPONDANCES, etc.)"),
    ("subCategory", "VARCHAR", "NULLABLE", "DMS sub-category"),
    ("allowedRoles", "Role[]", "DEFAULT []", "Roles permitted to access this document"),
    ("type", "ENUM(DocumentType)", "DEFAULT OTHER", "Legal document type classification"),
    ("status", "ENUM(DocumentStatus)", "DEFAULT UPLOADED", "Processing status (UPLOADED, INDEXED, etc.)"),
    ("uploaderId", "UUID", "FOREIGN KEY → User", "User who uploaded the file"),
    ("case_id", "UUID", "FK → Case, NULLABLE", "Associated case (if linked)"),
    ("isPending", "BOOLEAN", "DEFAULT false", "True if pending case association"),
    ("deletedAt", "TIMESTAMP", "NULLABLE", "Soft-delete timestamp"),
    ("createdAt", "TIMESTAMP", "DEFAULT NOW()", "Upload timestamp"),
]
t9 = doc.add_table(rows=1, cols=4)
t9.style = 'Table Grid'
for text, cell in zip(["Field", "Type", "Constraints", "Description"], t9.rows[0].cells):
    r = cell.paragraphs[0].add_run(text)
    set_font(r, 10, bold=True)
for row_data in dd_doc:
    row = t9.add_row().cells
    for text, cell in zip(row_data, row):
        r = cell.paragraphs[0].add_run(text)
        set_font(r, 10)
doc.add_paragraph()
page_break()

# ─────────────────────────────────────────────────────────────────────────────
# CHAPTER IV — IMPLEMENTATION, RESULTS AND TESTING
# ─────────────────────────────────────────────────────────────────────────────
chapter_title("CHAPTER IV", "IMPLEMENTATION, RESULTS AND TESTING")

section_heading("4.1", "Introduction")
para(
    "This chapter documents the implementation of each LexManage module, presents key code "
    "excerpts, and reports the results of unit, integration, and system testing. Screenshots "
    "of the running application are included for each major feature."
)

section_heading("4.2", "Implementation")

section_heading("4.2.1", "Authentication Module", size=12)
para(
    "The authentication module is implemented in "
    "lexmanage-backend/src/modules/auth/. It provides four REST endpoints: "
    "POST /auth/register, POST /auth/login, POST /auth/refresh, and POST /auth/logout."
)
para(
    "Registration creates a new Tenant record (for new firms) or validates an invitation token "
    "(for joiners). The tenantContext.runUnscoped() wrapper bypasses the row-level filter "
    "to allow cross-tenant email uniqueness checks before the user's tenant is known. "
    "Passwords are hashed with bcrypt at cost factor 12. JWT access tokens expire in 15 minutes; "
    "refresh tokens in 7 days. The refresh token is stored as a hash in the users table and "
    "rotated on each refresh call."
)
screenshot_box("Authentication — Registration Form")
screenshot_box("Authentication — Login Form")
screenshot_box("Authentication — Invitation Token Flow")

section_heading("4.2.2", "Case Management Module", size=12)
para(
    "Cases are managed in lexmanage-backend/src/modules/cases/. The CasesService "
    "implements full CRUD with Zod DTO validation. Cases support five statuses "
    "(OPEN, IN_PROGRESS, PENDING, CLOSED, ARCHIVED) and three priority levels. "
    "The assignee field links to a User record within the same tenant."
)
para(
    "On the frontend, the CaseManagementView renders a filterable, sortable table of cases. "
    "The NewCaseDialog component uses a multi-step form with client search autocomplete. "
    "The CaseDrawer side panel displays full case details, linked documents, deadlines, "
    "and a chat thread with LexAssist AI scoped to the case."
)
screenshot_box("Case Management — Case Table with Status Filters")
screenshot_box("Case Management — New Case Dialog")
screenshot_box("Case Management — Case Drawer (Detail View)")

section_heading("4.2.3", "Document Management System", size=12)
para(
    "The DMS is implemented in lexmanage-backend/src/modules/documents/. Document upload "
    "uses Multer's memory storage interceptor. The service pipeline is:"
)
bullet("1. Buffer file in memory (Multer FileInterceptor).")
bullet("2. Call fileTypeFromBuffer() to detect MIME type from magic bytes.")
bullet("3. Apply text/plain fallback for TXT files (no magic bytes).")
bullet("4. Validate MIME against allowedMimes list [application/pdf, .docx, text/plain].")
bullet("5. Construct MinIO object key: documents/{year}/{month}/{uuid8}/{filename}.")
bullet("6. Upload to tenant-scoped MinIO bucket (lex-{tenantId}) via MinioService.uploadFile().")
bullet("7. Create Document record in PostgreSQL (tenantId auto-injected by Prisma extension).")
bullet("8. Fire-and-forget n8n ingest webhook for PDF/DOCX files.")
bullet("9. Generate and return presigned download URL (1-hour TTL).")
para(
    "The 'Import to LexAssist AI' button calls POST /api/v1/ai/ingest-document with the "
    "document ID. The AiController verifies ownership, fetches the file buffer from MinIO, "
    "and calls N8nRagService.ingestDocument()."
)
screenshot_box("DMS — Document List with Category Filters")
screenshot_box("DMS — Upload Dropzone (PDF, DOCX, TXT Support)")
screenshot_box("DMS — Import to LexAssist AI Button")
screenshot_box("DMS — Document Detail / Download URL")

section_heading("4.2.4", "LexAssist AI Module", size=12)
para(
    "LexAssist AI is implemented as a bridge service (N8nRagService) that calls an external "
    "n8n workflow via HTTP webhooks. The n8n workflow implements the following RAG pipeline:"
)
bullet("1. Webhook Trigger receives: tenantId, userId, chatInput / fileData, sessionId.")
bullet("2. Document Loader decodes base64 file content (ingest) or passes query (chat).")
bullet("3. Cohere embed-english-v3.0 generates dense vector embeddings.")
bullet("4. Pinecone Vector Store upserts (ingest) or queries (chat) using tenantId as namespace.")
bullet("5. Cohere Reranker reranks top-K Pinecone results by relevance score.")
bullet("6. Simple Memory node loads/saves conversation history keyed by tenantId_userId_sessionId.")
bullet("7. Google Gemini 2.5 Flash LLM synthesises a grounded, cited answer.")
bullet("8. Webhook response returns the answer to N8nRagService, which returns it to the React frontend.")
para(
    "The ChatConversation and ChatMessage tables persist conversation history in PostgreSQL, "
    "ensuring history is preserved across browser sessions."
)
screenshot_box("LexAssist AI — Chat Interface with Conversation History")
screenshot_box("LexAssist AI — AI Response with Document Citation")
screenshot_box("LexAssist AI — New Conversation / Prompt Suggestions")
screenshot_box("n8n — Legal RAG Workflow Canvas")

section_heading("4.2.5", "Notifications Module", size=12)
para(
    "Notifications are managed in lexmanage-backend/src/modules/notifications/. "
    "The module uses two BullMQ queues ('mail' and 'reminders') backed by Redis. "
    "The NotificationsService exposes methods for creating, scheduling, and broadcasting "
    "notifications. Real-time delivery is handled by the EventsGateway (Socket.io WebSocket "
    "gateway), which maintains two room types: tenant_{tenantId} (firm-wide) and "
    "user_{userId} (private). The RemindersProcessor scans deadlines due within 3 days "
    "every 15 minutes via a NestJS @Cron decorator."
)
screenshot_box("Notifications — Notification Centre View")
screenshot_box("Notifications — Real-Time Bell Badge")
screenshot_box("Notifications — Send Notification Dialog")

section_heading("4.2.6", "Calendar Module", size=12)
para(
    "The calendar module provides a monthly view with deadline and event overlays. "
    "Deadlines are fetched from the /api/v1/calendar endpoint which returns events "
    "and deadlines for the queried month. Users can create, update, and mark deadlines "
    "as complete. Overdue deadlines are highlighted in red; upcoming ones in amber."
)
screenshot_box("Calendar — Monthly View with Deadline Overlays")
screenshot_box("Calendar — New Event / Deadline Dialog")

section_heading("4.2.7", "Client Management Module", size=12)
para(
    "The Clients module provides a searchable, paginated directory of firm clients. "
    "Each client has a name, email, phone, address, and type_client field. "
    "React Query hooks (useClients, useCreateClient, useUpdateClient, useDeleteClient) "
    "manage server state with automatic cache invalidation on mutations. "
    "Clients can be associated with one or more cases."
)
screenshot_box("Clients — Directory View with Search")
screenshot_box("Clients — Client Detail / Edit Modal")

section_heading("4.2.8", "Statistics Module", size=12)
para(
    "The StatsService assembles dashboard KPIs using 14 parallel Prisma queries executed "
    "via Promise.all() for optimal performance. Raw SQL is used only for the weekly "
    "time-series aggregations (date_trunc). The AI Dashboard endpoint (getAiDashboardData) "
    "returns counts of indexed documents, case summaries, and at-risk cases."
)
screenshot_box("Dashboard — KPI Cards with Month-Over-Month Deltas")
screenshot_box("Dashboard — Case Status Distribution Chart")
screenshot_box("Dashboard — Lawyer Workload Chart")
screenshot_box("Dashboard — 8-Week Activity Trend")
screenshot_box("AI Dashboard — Metrics and Case Summaries")

section_heading("4.3", "Testing")
para(
    "Three levels of testing were performed to validate correctness, integration, and "
    "end-to-end behaviour of the LexManage platform."
)

section_heading("4.3.1", "Unit Testing", size=12)
para(
    "Unit tests were written using Jest (NestJS's default test runner). Key unit tests include:"
)
bullet("TenantMiddleware: verifies that valid JWT sets tenantId in AsyncLocalStorage; expired token leaves tenantId undefined.")
bullet("AuthService.register(): verifies ConflictException on duplicate email; verifies new tenant creation on fresh registration.")
bullet("DocumentsService.upload(): verifies BadRequestException for unsupported MIME types; verifies text/plain acceptance; verifies MinIO object key format.")
bullet("N8nRagService.chat(): mocks fetch; verifies correct payload structure; verifies graceful fallback on network error.")
bullet("NotificationsService: verifies Socket.io emitToTenant call on notification creation.")
screenshot_box("Unit Test Results — Jest Test Suite Output")

section_heading("4.3.2", "Integration Testing", size=12)
para(
    "Integration tests validated the interaction between the NestJS controllers, "
    "services, and the PostgreSQL database (using a dedicated test database). Key tests:"
)
bullet("POST /auth/register → 201 Created + JWT tokens returned.")
bullet("POST /auth/login → 200 OK + tokens; 401 on wrong password.")
bullet("POST /documents/upload (PDF) → 201 + document object with presigned URL.")
bullet("POST /documents/upload (TXT) → 201 + document object with file_type: text/plain.")
bullet("POST /documents/upload (EXE) → 400 Bad Request (invalid MIME type).")
bullet("POST /ai/ingest-document (valid documentId) → 200 + {success: true}.")
bullet("POST /ai/dashboard-chat → 200 + {text: '...', sources: [...]}.")
bullet("Multi-tenant isolation: GET /cases with Tenant A token returns only Tenant A cases.")
screenshot_box("Integration Test Results — Supertest HTTP Test Suite")

section_heading("4.3.3", "System Testing", size=12)
para(
    "System testing validated the end-to-end user flows in the running application. "
    "Test scenarios executed:"
)
bullet("E2E-01: New firm registration → login → dashboard KPIs displayed correctly.")
bullet("E2E-02: Create case → add deadline → receive automated reminder notification.")
bullet("E2E-03: Upload TXT document → click 'Import to LexAssist AI' → ask question → receive document-grounded answer.")
bullet("E2E-04: Upload PDF → automatic n8n ingestion triggered → LexAssist answers questions about PDF content.")
bullet("E2E-05: Admin sends URGENT notification → all firm members receive real-time push notification.")
bullet("E2E-06: Language switch EN→FR → all UI labels update without page reload.")
bullet("E2E-07: Dark mode toggle → entire application switches theme.")
bullet("E2E-08: Invite colleague → colleague registers via token → joins the same tenant; cannot see another tenant's data.")
para("All eight E2E scenarios passed successfully in the final system test run.")
screenshot_box("System Test — End-to-End Test Plan and Results Matrix")
page_break()

# ─────────────────────────────────────────────────────────────────────────────
# CHAPTER V — CONCLUSIONS AND RECOMMENDATIONS
# ─────────────────────────────────────────────────────────────────────────────
chapter_title("CHAPTER V", "CONCLUSIONS AND RECOMMENDATIONS")

section_heading("5.1", "Conclusions")
para(
    "This project successfully designed, developed, and tested LexManage — a multi-tenant "
    "SaaS Legal Management Platform with an integrated AI-powered legal research assistant. "
    "The following objectives were met:"
)
bullet("A secure, JWT-based multi-tenant authentication system with five RBAC roles was implemented, with tenant isolation enforced at the ORM layer via Prisma query extensions.")
bullet("A comprehensive Case Management module was delivered, supporting full case lifecycle management (creation, assignment, status transitions, deadline tracking, document association).")
bullet("A Document Management System was built supporting PDF, DOCX, and TXT files, with tenant-scoped MinIO storage, role-based access control, and soft deletion.")
bullet("LexAssist AI was integrated as a RAG pipeline (Cohere + Pinecone + Gemini, orchestrated by n8n) with per-tenant vector namespace isolation and per-user conversation memory.")
bullet("A real-time notification system was implemented using Socket.io WebSocket rooms with firm-wide and user-private channels, complemented by BullMQ async email delivery.")
bullet("A calendar and deadline tracking module was built with automated cron-based 3-day-advance reminders.")
bullet("A bilingual (English/French) interface with dark/light theme support was delivered.")
para(
    "The system demonstrated that a modern, cloud-native legal platform can be built entirely "
    "on open-source components (NestJS, React, Prisma, PostgreSQL, MinIO, Redis, Socket.io) "
    "at minimal cost, while matching the feature depth of commercially licensed products. "
    "The Pinecone-namespaced multi-tenant RAG architecture proved effective: LexAssist "
    "correctly retrieved clause-level content from uploaded TXT, PDF, and DOCX documents "
    "and produced accurate, document-grounded answers during all system tests."
)

section_heading("5.2", "Recommendations")
para(
    "Based on the development experience and test results, the following recommendations "
    "are made for the system's operation and further development:"
)
bullet("Upgrade from Gemini free tier to a paid Gemini API plan (or integrate an OpenAI GPT-4 API) to eliminate rate-limiting constraints on LexAssist AI in production environments.")
bullet("Implement SSL/TLS termination and configure CORS origin whitelist before exposing the system to the public internet.")
bullet("Enable MinIO versioning and object-lock for document integrity compliance with archiving regulations.")
bullet("Implement a proper pagination cursor pattern on the DMS and Case Management endpoints (replacing offset-based pagination) to maintain performance at scale.")
bullet("Deploy using Docker Compose or Kubernetes with environment-specific secrets management (HashiCorp Vault or AWS Secrets Manager) rather than a plain .env file.")
bullet("Integrate a full-text search engine (Elasticsearch or PostgreSQL full-text search) alongside the vector search for keyword-based document queries.")

section_heading("5.3", "Perspectives for Further Study")
para("Several promising directions exist for future development of LexManage:")
bullet("Native Mobile Application: Develop React Native or Flutter clients for iOS and Android, extending LexAssist AI access to mobile devices for courtroom use.")
bullet("Multi-Language RAG Support: Extend LexAssist AI to support French-language embeddings (Cohere embed-multilingual-v3.0) for French-language legal documents and jurisdictions.")
bullet("Document Generation: Integrate AI-powered contract drafting using LLM function-calling to generate DOCX templates from case data.")
bullet("Billing and Invoicing Module: Implement time-tracking, billable hours recording, and invoice generation in PDF format.")
bullet("Advanced Analytics: Build a predictive analytics module that estimates case resolution time and success probability using historical case data.")
bullet("Offline-First Architecture: Implement service workers and IndexedDB for offline access to case summaries and documents when court connectivity is unavailable.")
bullet("Federated Multi-Firm Collaboration: Allow cross-tenant document sharing for joint case handling (e.g., co-counsel arrangements), with granular permission scoping.")
page_break()

# ─────────────────────────────────────────────────────────────────────────────
# REFERENCES
# ─────────────────────────────────────────────────────────────────────────────
heading("REFERENCES", level=0)
para("")
refs = [
    "Fowler, M. (2018). Patterns of Enterprise Application Architecture. Addison-Wesley.",
    "Richardson, C. (2019). Microservices Patterns: With Examples in Java. Manning Publications.",
    "Kleppmann, M. (2017). Designing Data-Intensive Applications. O'Reilly Media.",
    "NestJS Documentation. (2024). A progressive Node.js framework for building efficient and scalable server-side applications. Retrieved from https://docs.nestjs.com",
    "Prisma Documentation. (2024). Next-generation Node.js and TypeScript ORM. Retrieved from https://www.prisma.io/docs",
    "Pinecone Documentation. (2024). Vector database for machine learning applications. Retrieved from https://docs.pinecone.io",
    "Cohere Documentation. (2024). Cohere embeddings and reranking models. Retrieved from https://docs.cohere.com",
    "Google AI. (2024). Gemini API documentation. Retrieved from https://ai.google.dev/docs",
    "n8n Documentation. (2024). Workflow automation for technical teams. Retrieved from https://docs.n8n.io",
    "Lewis, P., Perez, E., Piktus, A., et al. (2020). Retrieval-Augmented Generation for Knowledge-Intensive NLP Tasks. Advances in Neural Information Processing Systems, 33.",
    "Grand View Research. (2023). Legal Technology Market Size, Share & Trends Analysis Report By Technology, By Deployment, By Application, By End-use, By Region, and Segment Forecasts, 2023–2030.",
    "OWASP Foundation. (2023). OWASP Top Ten — Ten Most Critical Web Application Security Risks. Retrieved from https://owasp.org/www-project-top-ten/",
    "React Documentation. (2024). React — The library for web and native user interfaces. Retrieved from https://react.dev",
    "Socket.IO Documentation. (2024). Bidirectional and low-latency communication for every platform. Retrieved from https://socket.io/docs",
    "MinIO Documentation. (2024). MinIO High Performance Object Storage. Retrieved from https://min.io/docs",
    "BullMQ Documentation. (2024). Premium message queue for Node.js based on Redis. Retrieved from https://docs.bullmq.io",
    "Tannenbaum, A. S., & Bos, H. (2014). Modern Operating Systems (4th ed.). Pearson Education.",
    "Sommerville, I. (2015). Software Engineering (10th ed.). Pearson Education.",
    "Bass, L., Clements, P., & Kazman, R. (2021). Software Architecture in Practice (4th ed.). Addison-Wesley.",
    "Zustand Documentation. (2024). Bear necessities for state management in React. Retrieved from https://docs.pmnd.rs/zustand",
]
for i, ref in enumerate(refs, 1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.first_line_indent = Cm(-0.5)
    p.paragraph_format.left_indent = Cm(0.5)
    r = p.add_run(f"[{i}] {ref}")
    set_font(r, 12)
page_break()

# ─────────────────────────────────────────────────────────────────────────────
# APPENDICES
# ─────────────────────────────────────────────────────────────────────────────
heading("APPENDICES", level=0)
para("")

section_heading("Appendix 1", "Environment Configuration (.env)")
para("The following environment variables must be set in lexmanage-backend/.env:")

env_vars = """DATABASE_URL="postgresql://postgres:password@localhost:5432/lexmanage"
JWT_SECRET="your-256-bit-secret-key"
JWT_ACCESS_EXPIRY="15m"
JWT_REFRESH_EXPIRY="7d"
MINIO_ENDPOINT="localhost"
MINIO_PORT=9000
MINIO_USE_SSL=false
MINIO_ACCESS_KEY="minioadmin"
MINIO_SECRET_KEY="minioadmin"
REDIS_HOST="localhost"
REDIS_PORT=6379
ALLOWED_ORIGINS="http://localhost:5173"
N8N_RAG_CHAT_URL="https://your-n8n-instance.app.n8n.cloud/webhook/legal-rag-chat"
N8N_RAG_INGEST_URL="https://your-n8n-instance.app.n8n.cloud/webhook/legal-rag-ingest"
MAIL_HOST="smtp.example.com"
MAIL_PORT=587
MAIL_USER="your-email@example.com"
MAIL_PASS="your-email-password"
"""
p = doc.add_paragraph()
r = p.add_run(env_vars)
set_font(r, 10)
r.font.name = "Courier New"

section_heading("Appendix 2", "API Endpoint Reference")
para("Table 10: Complete API Endpoint Reference")
api_data = [
    ("POST", "/api/v1/auth/register", "PUBLIC", "Register new firm + admin user"),
    ("POST", "/api/v1/auth/login", "PUBLIC", "Authenticate user, receive JWT"),
    ("POST", "/api/v1/auth/refresh", "PUBLIC", "Rotate refresh token"),
    ("POST", "/api/v1/auth/logout", "AUTH", "Invalidate refresh token"),
    ("GET",  "/api/v1/auth/me", "AUTH", "Get current user profile"),
    ("PATCH","/api/v1/auth/profile", "AUTH", "Update user profile"),
    ("GET",  "/api/v1/cases", "AUTH", "List all cases (paginated)"),
    ("POST", "/api/v1/cases", "LAWYER+", "Create new case"),
    ("GET",  "/api/v1/cases/:id", "AUTH", "Get case details"),
    ("PATCH","/api/v1/cases/:id", "LAWYER+", "Update case"),
    ("DELETE","/api/v1/cases/:id", "ADMIN", "Archive/delete case"),
    ("GET",  "/api/v1/clients", "AUTH", "List all clients"),
    ("POST", "/api/v1/clients", "AUTH", "Create client"),
    ("PATCH","/api/v1/clients/:id", "AUTH", "Update client"),
    ("DELETE","/api/v1/clients/:id", "AUTH", "Delete client"),
    ("POST", "/api/v1/documents/upload", "LAWYER+", "Upload document (multipart/form-data)"),
    ("GET",  "/api/v1/documents", "AUTH", "List documents (paginated, filterable)"),
    ("GET",  "/api/v1/documents/:id/download-url", "AUTH", "Get presigned download URL"),
    ("DELETE","/api/v1/documents/:id", "ADMIN", "Delete document"),
    ("POST", "/api/v1/ai/chat", "AUTH", "LexAssist chat (case context)"),
    ("POST", "/api/v1/ai/dashboard-chat", "AUTH", "LexAssist chat (dashboard)"),
    ("POST", "/api/v1/ai/ingest-document", "AUTH", "Ingest document to RAG"),
    ("GET",  "/api/v1/notifications", "AUTH", "List notifications"),
    ("POST", "/api/v1/notifications", "ADMIN", "Create/broadcast notification"),
    ("PATCH","/api/v1/notifications/:id/read", "AUTH", "Mark notification as read"),
    ("GET",  "/api/v1/stats/dashboard", "AUTH", "Dashboard KPI statistics"),
    ("GET",  "/api/v1/stats/ai-dashboard", "AUTH", "AI dashboard metrics"),
    ("GET",  "/api/v1/calendar", "AUTH", "Get calendar events and deadlines"),
    ("GET",  "/api/v1/search", "AUTH", "Global full-text search"),
    ("GET",  "/api/v1/audit", "ADMIN", "View audit logs"),
]
t10 = doc.add_table(rows=1, cols=4)
t10.style = 'Table Grid'
for text, cell in zip(["Method", "Endpoint", "Auth", "Description"], t10.rows[0].cells):
    r = cell.paragraphs[0].add_run(text)
    set_font(r, 10, bold=True)
for method, endpoint, auth, desc in api_data:
    row = t10.add_row().cells
    for text, cell in zip([method, endpoint, auth, desc], row):
        r = cell.paragraphs[0].add_run(text)
        set_font(r, 10)

# ─────────────────────────────────────────────────────────────────────────────
# Save
# ─────────────────────────────────────────────────────────────────────────────
output_path = r"C:\Users\hp\Desktop\LexManage_BTech_SWE_Report.docx"
doc.save(output_path)
print(f"Report saved to: {output_path}")
print(f"Pages estimated: ~{len(doc.paragraphs) // 30 + 1}")
